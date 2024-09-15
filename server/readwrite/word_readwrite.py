# -*- coding: utf-8 -*-
"""
    Word read / write
    Author: awtestergit
"""

# microsoft docx
from docx import Document
from interface.interface_readwrite import IDocReaderWriter
from interface.interface_model import llm_continue, ContinueExit

class WordReaderWriter(IDocReaderWriter):
    TYPE = "docx"

    def __init__(self) -> None:
        super().__init__()
        self.type = self.TYPE

    def read_doc_to_texts_by_page(self, doc_path:str|bytes, start_page:int=0, end_page:int=-1, remove_mark=[], streaming=True, continue_flag:llm_continue=None):
        """
        read texts page by page
        doc_path: full path of the document to read
        remove_mark: the marks (texts) to be removed, note: the whole line will be removed at mark
        streaming: if to stream by yield a generator
        output: a generator regardless of streaming
            if streaming, yield to page by page
            if not streaming, yield to list of texts from the document, each entry is texts in a page
        """
        output = []

        valid = self.__validate_doc_type__(doc_path=doc_path)
        # if not valid
        if not valid: # either file does not exist, or file extension is not 'pdf'
            raise ValueError(f"DocReaderWriter read document failed. either file is not {self.type}, or it does not exist! File: {doc_path}")
        else:
            #else, read docx
            # convert page number to index
            start_page = start_page -1 if start_page > 0 else 0
            end_page = end_page + 1 if end_page > 0 else -1 # to include reading in end_page, else -1
            check_start, check_end = True, True # flag to check start & end
            # if bytes, use bytesIO
            #with io.BytesIO(doc_path) if type(doc_path) is bytes else open(doc_path, 'rb') as file:
            result = []
            softbreak = False #track softbreak, for softbreak there's text in the run to be handled
            with open(doc_path, 'rb') as file: # read only
                current_page = 0
                document = Document(file)
                text = ''
                output = [] # page content holder
                for _, parag in enumerate(document.paragraphs):
                    #
                    # every parag is a paragraph
                    # parag.text is the paragraph, looks like [texts, softbreak, texts]
                    #   plus text in a soft pagebreak, i.e, parag.text in a softbreak belongs to the previous paragraph
                    #
                    # check continue
                    if continue_flag:
                        cf = continue_flag.check_continue_flag()
                        if not cf:
                            raise ContinueExit()

                    run_length = len(parag.runs)
                    for idx, run in enumerate(parag.runs):
                        # remove mark, if any
                        mark_removed = False
                        if len(remove_mark)>0:
                            for mark in remove_mark:
                                if run.text.find(mark)>-1:
                                    mark_removed = True
                                    break
                        if mark_removed: # do not include mark string line
                            continue

                        # check page break
                        softbreak = True if 'lastRenderedPageBreak' in run._element.xml else False
                        hardbreak = True if ('w:br' in run._element.xml and 'type="page"' in run._element.xml) else False
                        page_break = softbreak or hardbreak
                        # check start page
                        if check_start and start_page > 0:
                            if page_break:
                                # this is executed only when user set start page >= 2
                                if current_page == start_page - 1: # if current page is the same as one page before startpage
                                    output = [] # reset, empty all texts accumulated so far, because next page is the start page
                                    check_start = False # done with checking start
                        # check end page
                        if check_end and end_page > 0:
                            if page_break:
                                if current_page == end_page:
                                    check_end = False

                        #current_page = current_page + 1 if page_break else current_page # increment by 1
                        if page_break:
                            text = self.__remove_newline_from_text__(text)
                            text = text.strip()
                            output.append(text)
                            text = '\n'.join(t for t in output) # for each page
                            page_content = [text, f"page_{current_page}"]
                            result.append(page_content) # add to result

                            ## streaming ##
                            if streaming:
                                yield page_content
                            ##    
                            current_page = current_page + 1 if page_break else current_page # increment by 1
                            # reset
                            output = []
                            text = ''
                            if softbreak:
                                # treat softbreak text as a seperate paragraph, especially for pdf vs word comparation
                                text = run.text # this soft pagebreak text
                        else: # normal
                            text += run.text

                        if not check_end: # if so, meaning we have finished reading the end_page
                            break # break for run loop

                    # after run ends, remove the \n in the text
                    if len(text)>0:
                        text = self.__remove_newline_from_text__(text)
                        text = text.strip()
                        output.append(text)
                        text = '' #reset

                    if not check_end:
                        break # break for paragraph loop
                # check any leftover
                if len(output)>0:
                    text = '\n'.join(t for t in output)
                    page_content = [text, f"page_{current_page}"]
                    result.append(page_content) # add to result

                    ## streaming ##
                    if streaming:
                        yield page_content
                    ##    
        if not streaming: # if not streaming, yield whole result, which is a list
            yield from result

    def read_doc_to_texts_by_block(self, doc_path:str|bytes, start_page:int=0, end_page:int=-1, remove_mark=[], reconstruct_paragraph_at_pagebreak=True, strip=True, streaming=False, read_by=2, continue_flag:llm_continue=None):
        """
        'block' means paragraph - in word, it will be a paragraph, or a paragraph splitted by a soft pagebreak
                - in pdf, if a paragraph is splitted by a page, there is no way to recover it so far
        doc_path: full path of the document to read
        remove_mark: the marks (texts) to be removed, note: the whole line will be removed at mark
        reconstruct_paragraph_at_pagebreak: if True, will construct paragraph splitted by soft pagebreak in word, if False, the sub-paragraph splitted 
            by soft pagebreak will be treaded as a paragraph
            for pdf, it does nothing
        read_by: 0: page, 1: document, 2: paragraph, at this call, only 1 or 2
        output: a generator regardless of streaming
            if streaming, yield to paragraph by paragraph
            if not streaming, yield to list of texts from the document, each entry is texts in a paragraph
        """

        output = []

        valid = self.__validate_doc_type__(doc_path=doc_path)
        # if not valid
        if not valid: # either file does not exist, or file extension is not 'pdf'
            raise ValueError(f"DocReaderWriter read document failed. either file is not {self.type}, or it does not exist! File: {doc_path}")
        else:
            #else, read docx
            # convert page number to index
            start_page = start_page -1 if start_page > 0 else 0
            end_page = end_page + 1 if end_page > 0 else -1 # to include reading in end_page, else -1
            check_start, check_end = True, True # flag to check start & end
            # if bytes, use bytesIO
            #with io.BytesIO(doc_path) if type(doc_path) is bytes else open(doc_path, 'rb') as file:
            with open(doc_path, 'rb') as file: # read only
                current_page = 0
                document = Document(file)
                text = ''
                for _, parag in enumerate(document.paragraphs):
                    #
                    # parag.text is the paragraph, looks like [texts, softbreak, texts]
                    #   plus text in a soft pagebreak, i.e, parag.text in a softbreak belongs to the previous paragraph
                    #
                    # check continue
                    if continue_flag:
                        cf = continue_flag.check_continue_flag()
                        if not cf:
                            raise ContinueExit()

                    text = ''
                    run_length = len(parag.runs)
                    for idx, run in enumerate(parag.runs):
                        # remove mark, if any
                        mark_removed = False
                        if len(remove_mark)>0:
                            for mark in remove_mark:
                                if run.text.find(mark)>-1:
                                    mark_removed = True
                                    break
                        if mark_removed: # do not include mark string line
                            continue

                        # check page break
                        softbreak = True if 'lastRenderedPageBreak' in run._element.xml else False
                        hardbreak = True if ('w:br' in run._element.xml and 'type="page"' in run._element.xml) else False

                        if softbreak:
                            if reconstruct_paragraph_at_pagebreak: # if to make a complete paragraph
                                text += run.text
                            else: # treat softbreak text as a seperate paragraph, especially for pdf vs word comparation
                                text = self.__remove_newline_from_text__(text)
                                text = text.strip() if strip else text
                                parag_content = [text, f"page_{current_page}"]
                                output.append(parag_content)
                                ## stream ##
                                if streaming and read_by == 2: # if streaming and read by 'paragraph'
                                    yield parag_content

                                text = run.text # this soft pagebreak text
                        else: # normal
                            text += run.text

                        page_break = softbreak or hardbreak
                        # check start page
                        if check_start and start_page > 0:
                            if page_break:
                                # this is executed only when user set start page >= 2
                                if current_page == start_page - 1: # if current page is the same as one page before startpage
                                    output = [] # reset, empty all texts accumulated so far, because next page is the start page
                                    check_start = False # done with checking start
                        # check end page
                        if check_end and end_page > 0:
                            if page_break:
                                if current_page == end_page:
                                    check_end = False

                        #current_page = current_page + 1 if page_break else current_page # increment by 1
                        if page_break:
                            current_page = current_page + 1 if page_break else current_page # increment by 1

                        if not check_end: # if so, meaning we have finished reading the end_page
                            break # break for run loop

                    # after run ends, remove the \n in the text
                    if len(text)>0:
                        text = self.__remove_newline_from_text__(text)
                        text = text.strip() if strip else text
                        parag_content = [text, f"page_{current_page}"]
                        output.append(parag_content)

                        ## stream ##
                        if streaming and read_by == 2: # if streaming and read by 'paragraph'
                            yield parag_content

                    if not check_end:
                        break # break for paragraph loop

        if not streaming or read_by == 1: # if not streaming, or read by 'document'
            yield from output

    def write_text_to_doc(self, texts: str | list[str] | list[list[str]], output_path: str, template_path: str = None, **kwargs) -> bool:
        return super().write_text_to_doc(texts, output_path, template_path, **kwargs)